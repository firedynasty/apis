#!/usr/bin/env python3
"""
onenote.py - a tiny OneNote CLI for macOS using Microsoft Graph.

No pip installs needed (standard library only).

First-time setup:
  1. Put your Application (client) ID in CLIENT_ID below (or set ONENOTE_CLIENT_ID env var).
  2. Run:  python3 onenote.py login
     It prints a code + URL. Open the URL, paste the code, sign in with the
     Microsoft account that OWNS your notes. The token is cached locally.

Commands:
  login                         Sign in (device-code flow). Caches token.
  notebooks                     List your notebooks.
  sections [notebook-id]        List sections (optionally within one notebook).
  pages [section-id]            List pages (optionally within one section).
  download <page-id> [outfile]  Download a page's HTML to a file (or stdout).
  upload <file> [--title T] [--section SECTION_ID]
                                Upload a .html/.txt/.md file as a new page.

Token cache lives at ~/.onenote_token.json
"""

import json
import os
import sys
import time
import urllib.request
import urllib.parse
import urllib.error

# ---- CONFIG ----------------------------------------------------------------
CLIENT_ID = os.environ.get("ONENOTE_CLIENT_ID", "")
# 'common' accepts both personal and work/school accounts.
TENANT = "common"
SCOPES = "Notes.ReadWrite offline_access"
TOKEN_FILE = os.path.expanduser("~/.onenote_token.json")
GRAPH = "https://graph.microsoft.com/v1.0"
AUTH = f"https://login.microsoftonline.com/{TENANT}/oauth2/v2.0"
# ----------------------------------------------------------------------------


def _post(url, data):
    body = urllib.parse.urlencode(data).encode()
    req = urllib.request.Request(url, data=body, method="POST")
    req.add_header("Content-Type", "application/x-www-form-urlencoded")
    try:
        with urllib.request.urlopen(req) as r:
            return json.loads(r.read().decode())
    except urllib.error.HTTPError as e:
        return json.loads(e.read().decode())


def _graph(method, path, headers=None, data=None, raw=False):
    tok = _valid_token()
    url = path if path.startswith("http") else GRAPH + path
    for attempt in range(5):
        req = urllib.request.Request(url, data=data, method=method)
        req.add_header("Authorization", "Bearer " + tok)
        for k, v in (headers or {}).items():
            req.add_header(k, v)
        try:
            with urllib.request.urlopen(req, timeout=120) as r:
                content = r.read()
                if raw:
                    return content
                return json.loads(content.decode()) if content else {}
        except urllib.error.HTTPError as e:
            if e.code == 429:
                wait = int(e.headers.get("Retry-After", 5 * (attempt + 1)))
                print(f"Rate limited, retrying in {wait}s...")
                time.sleep(wait)
                continue
            sys.exit(f"Graph error {e.code}: {e.read().decode()}")
    sys.exit("Failed after 5 retries (rate limited)")


def _save(tokens):
    tokens["_expires_at"] = time.time() + int(tokens.get("expires_in", 3600)) - 60
    with open(TOKEN_FILE, "w") as f:
        json.dump(tokens, f)
    os.chmod(TOKEN_FILE, 0o600)


def _valid_token():
    if not os.path.exists(TOKEN_FILE):
        sys.exit("Not signed in. Run:  python3 onenote.py login")
    with open(TOKEN_FILE) as f:
        t = json.load(f)
    if time.time() < t.get("_expires_at", 0):
        return t["access_token"]
    # refresh
    if "refresh_token" in t:
        new = _post(AUTH + "/token", {
            "client_id": CLIENT_ID,
            "grant_type": "refresh_token",
            "refresh_token": t["refresh_token"],
            "scope": SCOPES,
        })
        if "access_token" in new:
            _save(new)
            return new["access_token"]
    sys.exit("Token expired and refresh failed. Run:  python3 onenote.py login")


def login():
    if CLIENT_ID == "PASTE_YOUR_CLIENT_ID_HERE":
        sys.exit("Set CLIENT_ID in the script (or ONENOTE_CLIENT_ID env var) first.")
    dc = _post(AUTH + "/devicecode", {"client_id": CLIENT_ID, "scope": SCOPES})
    if "user_code" not in dc:
        sys.exit("Device code request failed: " + json.dumps(dc, indent=2))
    print("\n  1. Open:", dc["verification_uri"])
    print("  2. Enter code:", dc["user_code"])
    print("  3. Sign in with the account that owns your notes.\n")
    interval = int(dc.get("interval", 5))
    while True:
        time.sleep(interval)
        r = _post(AUTH + "/token", {
            "client_id": CLIENT_ID,
            "grant_type": "urn:ietf:params:oauth:grant-type:device_code",
            "device_code": dc["device_code"],
        })
        if "access_token" in r:
            _save(r)
            print("Signed in. Token cached at", TOKEN_FILE)
            return
        if r.get("error") == "authorization_pending":
            continue
        if r.get("error") == "slow_down":
            interval += 5
            continue
        sys.exit("Login failed: " + json.dumps(r, indent=2))


def notebooks():
    for n in _graph("GET", "/me/onenote/notebooks").get("value", []):
        print(f'{n["id"]}\t{n["displayName"]}')


def sections(notebook_id=None):
    path = f"/me/onenote/notebooks/{notebook_id}/sections" if notebook_id else "/me/onenote/sections"
    for s in _graph("GET", path).get("value", []):
        print(f'{s["id"]}\t{s["displayName"]}')


def pages(section_id=None):
    path = f"/me/onenote/sections/{section_id}/pages" if section_id else "/me/onenote/pages"
    for p in _graph("GET", path).get("value", []):
        print(f'{p["id"]}\t{p.get("title", "(untitled)")}')


def download(page_id, outfile=None):
    html = _graph("GET", f"/me/onenote/pages/{page_id}/content", raw=True)
    if outfile:
        with open(outfile, "wb") as f:
            f.write(html)
        print("Saved", outfile)
    else:
        sys.stdout.write(html.decode("utf-8", "replace"))


def update(page_id, path):
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    body = ""
    for line in content.splitlines():
        line = line.strip()
        if line:
            escaped = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            body += f"<p>{escaped}</p>"
    patch = json.dumps([{"target": "body", "action": "replace", "content": body}])
    _graph("PATCH", f"/me/onenote/pages/{page_id}/content",
           headers={"Content-Type": "application/json"},
           data=patch.encode("utf-8"))
    print("Page updated.")


def upload(path, title=None, section_id=None):
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    if title is None:
        title = os.path.splitext(os.path.basename(path))[0]
    # If it's not already a full HTML doc, wrap it.
    if "<html" not in content.lower():
        body = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        body = "".join(f"<p>{line}</p>" for line in body.splitlines() if line.strip())
        content = f"<!DOCTYPE html><html><head><title>{title}</title></head><body>{body}</body></html>"
    endpoint = (f"/me/onenote/sections/{section_id}/pages" if section_id
                else "/me/onenote/pages")
    res = _graph("POST", endpoint, headers={"Content-Type": "text/html"},
                 data=content.encode("utf-8"))
    print("Created page:", res.get("id", "(unknown id)"))
    links = res.get("links", {})
    if links.get("oneNoteWebUrl", {}).get("href"):
        print("View:", links["oneNoteWebUrl"]["href"])


def upload_folder(folder, title=None, section_id=None):
    import mimetypes
    import importlib.util
    import re as _re
    import html as _html

    IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}

    # Load mdToHtml converter for proper markdown rendering
    _md_to_html = None
    _md_converter_path = "/Users/stanleytan/Documents/technical/python/convert_html_to_docx/02evernote/mdToHtml.py"
    if os.path.exists(_md_converter_path):
        spec = importlib.util.spec_from_file_location("mdToHtml", _md_converter_path)
        _mdmod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(_mdmod)
        _md_to_html = _mdmod.md_to_html

    if title is None:
        title = os.path.basename(folder.rstrip("/").rstrip("\\")) or "Untitled"

    # Skip hidden files (e.g. .DS_Store) and only include regular files
    files = sorted(f for f in os.listdir(folder)
                   if os.path.isfile(os.path.join(folder, f)) and not f.startswith("."))

    body_html = ""
    image_parts = []

    for fname in files:
        fpath = os.path.join(folder, fname)
        ext = os.path.splitext(fname)[1].lower()
        heading = os.path.splitext(fname)[0].replace("-", " ").replace("_", " ")

        if ext == ".txt":
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
            escaped = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            inner = "".join(f"<p>{line}</p>" for line in escaped.splitlines() if line.strip())
            body_html += f"<h2>{heading}</h2>{inner}"

        elif ext == ".md":
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
            if _md_to_html:
                inner = _md_to_html(content)
            else:
                escaped = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                inner = "".join(f"<p>{line}</p>" for line in escaped.splitlines() if line.strip())
            body_html += f"<h2>{heading}</h2>{inner}"

        elif ext in {".html", ".htm"}:
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
            m = _re.search(r'<body[^>]*>(.*?)</body>', content, _re.S | _re.I)
            inner = m.group(1) if m else content
            body_html += f"<h2>{heading}</h2>{inner}"

        elif ext == ".docx":
            import subprocess
            inner = None
            # 1. Pandoc — best quality, preserves bold/headers/tables/lists
            try:
                result = subprocess.run(
                    ["pandoc", fpath, "-t", "html", "--no-highlight"],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    inner = result.stdout
                else:
                    raise RuntimeError("pandoc produced no output")
            except subprocess.TimeoutExpired:
                print(f"SKIPPED: {fname} — pandoc timed out (file may be too large or corrupted)")
            except (FileNotFoundError, RuntimeError):
                # 2. python-docx — plain paragraphs, no formatting
                try:
                    import docx as _docx
                    doc = _docx.Document(fpath)
                    inner = "".join(
                        f"<p>{_html.escape(p.text)}</p>"
                        for p in doc.paragraphs if p.text.strip()
                    )
                except ImportError:
                    # 3. textutil — macOS built-in, plain text only
                    try:
                        result = subprocess.run(
                            ["textutil", "-convert", "txt", "-stdout", fpath],
                            capture_output=True, text=True
                        )
                        if result.returncode == 0 and result.stdout.strip():
                            escaped = result.stdout.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                            inner = "".join(f"<p>{line}</p>" for line in escaped.splitlines() if line.strip())
                        else:
                            print(f"SKIPPED: {fname} — all docx converters failed")
                    except Exception as e:
                        print(f"SKIPPED: {fname} — {e}")
                except Exception as e:
                    print(f"SKIPPED: {fname} — {e}")
            if inner:
                body_html += f"<h2>{heading}</h2>{inner}"

        elif ext in IMAGE_EXTS:
            mime = mimetypes.guess_type(fname)[0] or "image/png"
            safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in fname)
            body_html += f'<p><img src="name:{safe_name}" /></p>'
            with open(fpath, "rb") as f:
                data = f.read()
            image_parts.append((safe_name, mime, data))

        else:
            print(f"SKIPPED: {fname} — unsupported format")

    if not body_html:
        sys.exit("No supported files (.txt, .md, .html, .jpg, .png, etc.) found in folder.")

    html_doc = (
        f'<!DOCTYPE html><html><head><title>{title}</title></head>'
        f'<body>{body_html}</body></html>'
    )

    boundary = "OneNoteUploadBoundary"
    body = b""

    # HTML part — Content-Disposition name="Presentation" is required by OneNote API
    body += f'--{boundary}\r\nContent-Disposition: form-data; name="Presentation"\r\nContent-Type: text/html\r\n\r\n'.encode()
    body += html_doc.encode("utf-8")
    body += b"\r\n"

    # Image parts
    for safe_name, mime, data in image_parts:
        body += f'--{boundary}\r\nContent-Disposition: form-data; name="{safe_name}"\r\nContent-Type: {mime}\r\n\r\n'.encode()
        body += data
        body += b"\r\n"

    body += f"--{boundary}--\r\n".encode()

    endpoint = (f"/me/onenote/sections/{section_id}/pages" if section_id
                else "/me/onenote/pages")

    size_mb = len(body) / 1024 / 1024
    print(f"Uploading '{title}' ({len(image_parts)} image(s), {size_mb:.1f} MB) — may take a moment...")
    res = _graph("POST", endpoint,
                 headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
                 data=body)
    print("Created page:", res.get("id", "(unknown id)"))
    links = res.get("links", {})
    if links.get("oneNoteWebUrl", {}).get("href"):
        print("View:", links["oneNoteWebUrl"]["href"])


def upload_files(filepaths, title=None, section_id=None, max_mb=18):
    """Bundle a list of files into one OneNote page (chunked if too large)."""
    import mimetypes
    import importlib.util
    import re as _re
    import html as _html
    import subprocess

    IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}

    _md_to_html = None
    _md_converter_path = "/Users/stanleytan/Documents/technical/python/convert_html_to_docx/02evernote/mdToHtml.py"
    if os.path.exists(_md_converter_path):
        spec = importlib.util.spec_from_file_location("mdToHtml", _md_converter_path)
        _mdmod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(_mdmod)
        _md_to_html = _mdmod.md_to_html

    if title is None:
        title = os.path.splitext(os.path.basename(filepaths[0]))[0]

    MAX_BYTES = max_mb * 1024 * 1024
    chunks = []       # list of (body_html, image_parts)
    cur_html = ""
    cur_imgs = []
    cur_size = 0

    for fpath in filepaths:
        fname = os.path.basename(fpath)
        if fname.startswith("."):
            continue
        ext = os.path.splitext(fname)[1].lower()
        heading = os.path.splitext(fname)[0].replace("-", " ").replace("_", " ")
        html_snip = None
        img_parts = []

        print(f"  Processing {fname}...")

        if ext == ".txt":
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
            escaped = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            inner = "".join(f"<p>{line}</p>" for line in escaped.splitlines() if line.strip())
            html_snip = f"<h2>{heading}</h2>{inner}"

        elif ext == ".md":
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
            inner = _md_to_html(content) if _md_to_html else "".join(
                f"<p>{line}</p>" for line in content.splitlines() if line.strip())
            html_snip = f"<h2>{heading}</h2>{inner}"

        elif ext in {".html", ".htm"}:
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
            m = _re.search(r'<body[^>]*>(.*?)</body>', content, _re.S | _re.I)
            inner = m.group(1) if m else content
            html_snip = f"<h2>{heading}</h2>{inner}"

        elif ext == ".docx":
            inner = None
            try:
                result = subprocess.run(
                    ["pandoc", fpath, "-t", "html", "--no-highlight"],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    inner = result.stdout
                else:
                    raise RuntimeError("pandoc produced no output")
            except subprocess.TimeoutExpired:
                print(f"  SKIPPED: {fname} — pandoc timed out")
            except (FileNotFoundError, RuntimeError):
                try:
                    import docx as _docx
                    doc = _docx.Document(fpath)
                    inner = "".join(f"<p>{_html.escape(p.text)}</p>"
                                    for p in doc.paragraphs if p.text.strip())
                except ImportError:
                    try:
                        r = subprocess.run(["textutil", "-convert", "txt", "-stdout", fpath],
                                           capture_output=True, text=True)
                        if r.returncode == 0 and r.stdout.strip():
                            escaped = r.stdout.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                            inner = "".join(f"<p>{line}</p>" for line in escaped.splitlines() if line.strip())
                        else:
                            print(f"  SKIPPED: {fname} — all converters failed")
                    except Exception as e:
                        print(f"  SKIPPED: {fname} — {e}")
                except Exception as e:
                    print(f"  SKIPPED: {fname} — {e}")
            if inner:
                html_snip = f"<h2>{heading}</h2>{inner}"

        elif ext in IMAGE_EXTS:
            mime = mimetypes.guess_type(fname)[0] or "image/png"
            safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in fname)
            with open(fpath, "rb") as f:
                data = f.read()
            html_snip = f'<p><img src="name:{safe_name}" /></p>'
            img_parts = [(safe_name, mime, data)]

        else:
            print(f"  SKIPPED: {fname} — unsupported format")

        if not html_snip:
            continue

        snip_size = len(html_snip.encode("utf-8")) + sum(len(d) for _, _, d in img_parts)
        if cur_size + snip_size > MAX_BYTES and cur_html:
            chunks.append((cur_html, cur_imgs))
            cur_html, cur_imgs, cur_size = "", [], 0
        cur_html += html_snip
        cur_imgs += img_parts
        cur_size += snip_size

    if cur_html:
        chunks.append((cur_html, cur_imgs))

    if not chunks:
        sys.exit("No supported files could be processed.")

    total = len(chunks)
    endpoint = (f"/me/onenote/sections/{section_id}/pages" if section_id
                else "/me/onenote/pages")

    print(f"\nUploading as {total} page(s)...")
    for i, (body_html, image_parts) in enumerate(chunks, 1):
        page_title = f"{title} ({i})" if total > 1 else title
        boundary = "OneNoteUploadBoundary"
        html_doc = (f'<!DOCTYPE html><html><head><title>{page_title}</title></head>'
                    f'<body>{body_html}</body></html>')
        body = b""
        body += f'--{boundary}\r\nContent-Disposition: form-data; name="Presentation"\r\nContent-Type: text/html\r\n\r\n'.encode()
        body += html_doc.encode("utf-8")
        body += b"\r\n"
        for safe_name, mime, data in image_parts:
            body += f'--{boundary}\r\nContent-Disposition: form-data; name="{safe_name}"\r\nContent-Type: {mime}\r\n\r\n'.encode()
            body += data
            body += b"\r\n"
        body += f"--{boundary}--\r\n".encode()

        size_mb = len(body) / 1024 / 1024
        print(f"  [{i}/{total}] '{page_title}' ({len(image_parts)} image(s), {size_mb:.1f} MB)...")
        res = _graph("POST", endpoint,
                     headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
                     data=body)
        print(f"  Created: {res.get('id', '(unknown)')}")
        links = res.get("links", {})
        if links.get("oneNoteWebUrl", {}).get("href"):
            print(f"  View: {links['oneNoteWebUrl']['href']}")
        if i < total:
            time.sleep(2)

    print(f"\nDone. {total} page(s) uploaded.")


def main():
    args = sys.argv[1:]
    if not args:
        print(__doc__)
        return
    cmd, rest = args[0], args[1:]
    if cmd == "login":
        login()
    elif cmd == "notebooks":
        notebooks()
    elif cmd == "sections":
        sections(rest[0] if rest else None)
    elif cmd == "pages":
        pages(rest[0] if rest else None)
    elif cmd == "download":
        if not rest:
            sys.exit("Usage: download <page-id> [outfile]")
        download(rest[0], rest[1] if len(rest) > 1 else None)
    elif cmd == "update":
        if len(rest) < 2:
            sys.exit("Usage: update <page-id> <file>")
        update(rest[0], rest[1])
    elif cmd == "upload":
        if not rest:
            sys.exit("Usage: upload <file> [--title T] [--section SECTION_ID]")
        f = rest[0]
        title = None
        section = None
        i = 1
        while i < len(rest):
            if rest[i] == "--title" and i + 1 < len(rest):
                title = rest[i + 1]; i += 2
            elif rest[i] == "--section" and i + 1 < len(rest):
                section = rest[i + 1]; i += 2
            else:
                i += 1
        upload(f, title, section)
    elif cmd == "upload-folder":
        if not rest:
            sys.exit("Usage: upload-folder <folder> [--title T] [--section SECTION_ID]")
        folder = rest[0]
        title = None
        section = None
        i = 1
        while i < len(rest):
            if rest[i] == "--title" and i + 1 < len(rest):
                title = rest[i + 1]; i += 2
            elif rest[i] == "--section" and i + 1 < len(rest):
                section = rest[i + 1]; i += 2
            else:
                i += 1
        upload_folder(folder, title, section)
    elif cmd == "upload-files":
        if not rest:
            sys.exit("Usage: upload-files [--title T] [--section SECTION_ID] file1 file2 ...")
        title = None
        section = None
        files = []
        i = 0
        while i < len(rest):
            if rest[i] == "--title" and i + 1 < len(rest):
                title = rest[i + 1]; i += 2
            elif rest[i] == "--section" and i + 1 < len(rest):
                section = rest[i + 1]; i += 2
            else:
                files.append(rest[i]); i += 1
        if not files:
            sys.exit("No files specified.")
        upload_files(files, title, section)
    else:
        print(__doc__)


if __name__ == "__main__":
    main()

